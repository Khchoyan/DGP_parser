import platform
import pandas as pd
import re
import requests
from bs4 import BeautifulSoup as bs

import docx
import os

import time
import datetime
from calendar import monthrange

pd.set_option('display.max_rows', None)
pd.set_option('display.max_columns', None)


def str_month2digit_month(month):
    """
    Функция переводит название месяца в его номер.
    """
    month = month.strip().lower()
    if month == 'январь':
        return '01'
    elif month == 'январь-февраль':
        return '02'
    elif month == 'январь-март':
        return '03'
    elif month == 'январь-апрель':
        return '04'
    elif month == 'январь-май':
        return '05'
    elif month == 'январь-июнь':
        return '06'
    elif month == 'январь-июль':
        return '07'
    elif month == 'январь-август':
        return '08'
    elif month == 'январь-сентябрь':
        return '09'
    elif month == 'январь-октябрь':
        return '10'
    elif month == 'январь-ноябрь':
        return '11'
    elif month == 'январь-декабрь':
        return '12'
    else:
        return 'unknown'


def get_os_type():
    os_type = platform.system()
    if os_type == "Windows":
        return "Windows"
    elif os_type == "Darwin" or os_type == "Linux":
        return "Unix"
    else:
        return "Unknown"


def doc_to_docx(path: str):
    """
    Функция конвертирует документ формата .doc в формат .docx
    doc_path - абсолютный путь к документу
    """
    exist_system = get_os_type()
    if exist_system == 'Unix':
        import doc2docx
        doc2docx.convert(path)

    elif exist_system == 'Windows':
        from win32com import client as wc
        w = wc.Dispatch('Word.Application')
        # Or use the following method to start a separate process:
        # w = wc.DispatchEx('Word.Application')
        doc = w.Documents.Open(path)
        doc.SaveAs(path + 'x', 16)
        doc.Close()
        w.Quit()
        print(f'Document {path} was converted to docx-format.')

    return path + 'x'


def reformat_date(date: str):
    """
    Функция переформатирует даты
    """
    date = date.strip()
    if date == 'I квартал':
        date = '31 march'
    elif date == 'II квартал':
        date = '30 june'
    elif date == 'I полугодие':
        date = '30 june'
    elif date == 'III квартал':
        date = '30 september'
    elif date == 'Январь-сентябрь':
        date = '30 september'
    elif date == 'IV квартал':
        date = '31 december'
    elif date == 'Год':
        date = '31 december'
    return date


def pars_year_by_months(year):
    """
    Функция для получения ссылок на документы по месяцам.
    Для ВВП реализовано возвращение названия последнего доступного месяца в конкретном году
    и ссылки на соответствующий раздел.
    """
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }

    url = f'https://rosstat.gov.ru/storage/mediabank/Doklad_{year}.htm'
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links_1 = pd.DataFrame()
    for i in range(0, len(soup.find('table').find_all('tr')[1].find_all('tr')), 2):
        month_name = soup.find('table').find_all('tr')[1].find_all('tr')[i].find_all('td')[0].text
        month_name = month_name.replace('\n', '')
        if month_name.split()[-1].lower() == 'год':
            month_name = 'Январь-декабрь'
        dok_link = soup.find('table').find_all('tr')[1].find_all('tr')[i].find_all('td')[1].find_all('a')[0].get('href')
        if dok_link[:4] != 'http':
            dok_link = 'https://rosstat.gov.ru' + dok_link
        pril_link = soup.find('table').find_all('tr')[1].find_all('tr')[i + 1].find_all('td')[0].find_all('a')[0].get(
            'href')
        if pril_link[:4] != 'http':
            pril_link = 'https://rosstat.gov.ru' + pril_link
        links_1 = links_1._append([[month_name, dok_link, pril_link]])

    # ВВП обновлляется 4 раза в год
    # links_1 = links_1[links_1[0].str.contains('март|май|август|ноябрь')]
    links_1 = links_1[links_1[0].str.contains('май|август|ноябрь')]
    links_1.columns = ['Месяц', 'Ссылка', 'Дополнительная ссылка']
    links_1 = links_1.iloc[::-1].reset_index(drop=True)

    if links_1.empty:
        return 0, 0
    else:
        return links_1['Месяц'].iloc[-1], links_1['Ссылка'].iloc[-1]


def download_document(year, month, url):
    """
    Функция скачивает документ с данными по инвестициям за конкретный месяц.
    year - год в формате ХХХХ.
    month - полное название месяца на русском языке.
    url - ссылка на документ.
    Первые две переменные необходимы для назначения имени скачиваемому файлу.
    Возвращает путь к сохранённому файлу.
    """
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links = pd.DataFrame()
    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        dok_link = link.get('href')
        links = links._append([[branch_name, dok_link]])

    indicator = 'Объем ВВП'
    if len(links[links[0] == indicator][1]) == 0:
        print(f'NO DOCUMENT {year}_{month}: {indicator}')
    else:
        link_to_download = links[links[0] == indicator][1].values[0]
        dok_name_to_download = f'{year}_{month}-2-4-0.doc'  # 2024_02-2-4-0.doc
        folder = os.getcwd()
        folder = os.path.join(folder, 'word_data', dok_name_to_download)

        response = requests.get(link_to_download, headers=header)
        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def parse_docx_document(path, year, month):
    """
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    """
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0

    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for i, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: ' ' + str(x))
    data_table = data_table[data_table.iloc[:, 0].str.contains('I квартал|I полугодие|Январь-сентябрь|IV квартал|Год')]

    if data_table.empty:
        data_table = [[] for _ in range(len(doc.tables[1].rows))]
        for i, row in enumerate(doc.tables[1].rows):
            for cell in row.cells:
                data_table[i].append(cell.text)

    for i in [1, 2, ]:
        data_table.iloc[:, i] = data_table.iloc[:, i].str.replace(' ', '').str.replace('\xa0', '').str.replace(',', '.')
    print(f'Document {path} was parsed')

    df_1 = data_table[[0, 1]]
    df_2 = data_table[[0, 2]]
    df_1 = df_1[df_1.iloc[:, 0].str.contains(' I квартал|I полугодие|Январь-сентябрь|Год')]
    df_2 = df_2[df_2.iloc[:, 0].str.contains(' I квартал|II квартал|III квартал|IV квартал')]

    df_1.iloc[:, 0] = df_1.iloc[:, 0].apply(lambda x: reformat_date(x))
    df_2.iloc[:, 0] = df_2.iloc[:, 0].apply(lambda x: reformat_date(x))

    for i in range(len(df_1)):
        if i < 4:
            df_1.iloc[i, 0] = pd.to_datetime(df_1.iloc[i, 0] + str(year - 1))
            df_2.iloc[i, 0] = pd.to_datetime(df_2.iloc[i, 0] + str(year - 1))
        else:
            df_1.iloc[i, 0] = pd.to_datetime(df_1.iloc[i, 0] + str(year))
            df_2.iloc[i, 0] = pd.to_datetime(df_2.iloc[i, 0] + str(year))

    return df_1, df_2


def create_new_date(last_date_in_file_year, last_date_in_file_month):
    now = datetime.datetime.now()
    lst_date = []
    _, last_day = monthrange(now.year, now.month)
    last_date = datetime.datetime.strptime(f"{now.year}-{now.month}-{last_day}", "%Y-%m-%d").date()

    for i in range((last_date.year - last_date_in_file_year) * 12 + last_date.month - last_date_in_file_month - 1):
        if last_date.month - 1 != 0:
            _, last_day = monthrange(last_date.year, last_date.month - 1)
            last_date = datetime.datetime.strptime(f"{last_date.year}-{last_date.month - 1}-{last_day}",
                                                   "%Y-%m-%d").date()
        else:
            _, last_day = monthrange(last_date.year - 1, 12)
            last_date = datetime.datetime.strptime(f"{last_date.year - 1}-{12}-{last_day}", "%Y-%m-%d").date()
        lst_date.append(last_date)
    return sorted(lst_date)


def append_date_rez_file_Y(xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет дабавление месяцев, если их нет в файле.
    """
    data_xlsx = pd.read_excel(xlsx_path)
    year = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).year
    month = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).month
    date_lst = create_new_date(year, month)
    for date in date_lst:
        new_string = {'Целевой показатель': [date]}
        new_string.update({c: [None] for c in data_xlsx.columns[1:]})
        new_string = pd.DataFrame(new_string)
        if not data_xlsx.empty and not new_string.empty:
            data_xlsx = pd.concat([data_xlsx, new_string])
    data_xlsx.to_excel(xlsx_path, index=False)


def update_rez_file_y(df_1, df_2, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    if df_1.values[-1][0] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = 'Реальный ВВП (Темп роста накопленным итогом)'
    name_2 = 'Реальный ВВП (Темп роста, квартал к кварталу прошлого года)'
    for j in df_1.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_1] = float(j[1].replace(',', '.'))
    for k in df_2.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == k[0], name_2] = float(k[1].replace(',', '.'))

    data_xlsx.to_excel(xlsx_path, index=False)


def main():
    """
    Основная функция. Выполняет проверку данных на полноту. Скачивет недостающие
    данные и дополняет ими файл с данными.
    """
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Реальный ВВП (Темп роста, '
                                                                                           'квартал к кварталу '
                                                                                           'прошлого года)']).iloc[
                                            -1]['Целевой показатель']).year
    if now - last_year_in_table < 2:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table + 1, now + 1):
            years.append(y)
    for year in years:
        time.sleep(15)
        month, URL = pars_year_by_months(year)
        if month == 0:
            print('Новых данных по ВВП не появилось')
            break
        time.sleep(60)
        path_to_docfile = download_document(year, month, URL)
        path = doc_to_docx(path_to_docfile)
        df_1, df_2 = parse_docx_document(path, year=year, month=month)
        os.remove(path_to_docfile)
        if not df_1.empty and not df_1.empty:
            update_rez_file_y(df_1, df_2)


if __name__ == '__main__':
    main()